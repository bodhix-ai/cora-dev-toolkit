"""
KB Processor Lambda - Asynchronous Document Processing

Triggered by: SQS messages from kb-document Lambda

Processing Pipeline:
1. Receive SQS message with document metadata
2. Download document from S3
3. Parse document (PDF, DOCX, TXT, MD)
4. Chunk text with overlap
5. Generate embeddings using configured AI provider
6. Store chunks with embeddings in kb_chunks table
7. Update document status (processing → indexed/failed)

No HTTP routes - SQS event handler only
"""

import json
import os
import re
import traceback
from typing import Dict, Any, Optional, List, Tuple
from io import BytesIO

import boto3
import org_common as common

# Environment variables
AWS_REGION = os.environ.get('AWS_REGION', 'us-east-1')

# AWS clients
s3_client = boto3.client('s3', region_name=AWS_REGION)
bedrock_runtime = boto3.client('bedrock-runtime', region_name=AWS_REGION)

# Processing constants
DEFAULT_CHUNK_SIZE = 1000  # characters
DEFAULT_CHUNK_OVERLAP = 200  # characters
DEFAULT_EMBEDDING_MODEL = 'amazon.titan-embed-text-v2:0'
DEFAULT_EMBEDDING_DIMENSION = 1024
MAX_RETRIES = 3
BATCH_SIZE = 25  # Bedrock limit for batch embeddings


def lambda_handler(event, context):
    """Main Lambda handler for SQS events."""
    try:
        # Process each SQS record
        for record in event.get('Records', []):
            process_sqs_message(record)
        
        return common.success_response({'message': 'Processing complete'})
    
    except Exception as e:
        print(f"Unhandled error in lambda_handler: {str(e)}")
        print(traceback.format_exc())
        return common.internal_error_response('Processing failed')



def process_sqs_message(record: Dict):
    """Process a single SQS message."""
    try:
        message_body = json.loads(record.get('body', '{}'))
        
        document_id = message_body.get('documentId')
        kb_id = message_body.get('kbId')
        s3_bucket = message_body.get('s3Bucket')
        s3_key = message_body.get('s3Key')
        action = message_body.get('action', 'index')
        
        if not all([document_id, kb_id, s3_bucket, s3_key]):
            print(f"Invalid message: missing required fields")
            return
        
        print(f"Processing document {document_id} from {s3_bucket}/{s3_key}")
        
        # Get document to extract org_id (CORA Compliance)
        document = common.find_one(
            table='kb_docs',
            filters={'id': document_id, 'is_deleted': False}
        )
        
        if not document:
            print(f"Document {document_id} not found")
            return
        
        org_id = document.get('org_id')
        if not org_id:
            print(f"Document {document_id} has no org_id")
            return
        
        # Update status to processing
        update_document_status(document_id, org_id, 'processing')
        
        # Process document
        if action == 'index':
            process_document(document_id, kb_id, s3_bucket, s3_key)
        else:
            print(f"Unknown action: {action}")
    
    except Exception as e:
        print(f"Error processing SQS message: {str(e)}")
        print(traceback.format_exc())
        # Let message go to DLQ for manual inspection


def process_document(document_id: str, kb_id: str, s3_bucket: str, s3_key: str):
    """Main document processing pipeline."""
    try:
        # Step 1: Get document and extract org_id for multi-tenancy validation (CORA Compliance)
        document = common.find_one(
            table='kb_docs',
            filters={'id': document_id, 'is_deleted': False}
        )
        
        if not document:
            raise ValueError(f"Document {document_id} not found")
        
        org_id = document.get('org_id')
        if not org_id:
            raise ValueError(f"Document {document_id} has no org_id")
        
        # Validate document belongs to the specified KB
        if document.get('kb_id') != kb_id:
            raise ValueError(f"Document {document_id} does not belong to KB {kb_id}")
        
        print(f"Processing document {document_id} for org {org_id}")
        
        # Step 2: Download document from S3
        print(f"Downloading document from S3: {s3_key}")
        document_bytes, mime_type = download_from_s3(s3_bucket, s3_key)
        
        # Step 3: Parse document based on MIME type
        print(f"Parsing document (MIME: {mime_type})")
        text, metadata = parse_document(document_bytes, mime_type, s3_key)
        
        if not text or len(text.strip()) < 10:
            raise ValueError("Document appears to be empty or too short")
        
        print(f"Extracted {len(text)} characters from document")
        
        # Step 4: Get embedding configuration (includes chunking parameters from ai_cfg_sys_rag)
        embedding_config = get_embedding_config()
        chunk_size = embedding_config.get('chunkSize', DEFAULT_CHUNK_SIZE)
        chunk_overlap = embedding_config.get('chunkOverlap', DEFAULT_CHUNK_OVERLAP)
        embedding_model = embedding_config.get('model', DEFAULT_EMBEDDING_MODEL)
        
        # Step 5: Chunk text
        print(f"Chunking text (size: {chunk_size}, overlap: {chunk_overlap})")
        chunks = chunk_text(text, chunk_size, chunk_overlap)
        
        print(f"Created {len(chunks)} chunks")
        
        # Step 6: Generate embeddings
        print(f"Generating embeddings with model: {embedding_model}")
        embeddings = generate_embeddings(chunks, embedding_model)
        
        if len(embeddings) != len(chunks):
            raise ValueError(f"Embedding count mismatch: {len(embeddings)} != {len(chunks)}")
        
        # Step 8: Store chunks with embeddings (with org_id for multi-tenancy)
        print(f"Storing {len(chunks)} chunks in database")
        store_chunks(document_id, kb_id, org_id, chunks, embeddings, embedding_model)
        
        # Step 9: Update document metadata (with org_id filter)
        update_document_metadata(document_id, org_id, metadata, len(chunks))
        
        # Step 10: Update status to indexed (with org_id filter)
        update_document_status(document_id, org_id, 'indexed')
        
        print(f"Successfully processed document {document_id}")
    
    except Exception as e:
        error_message = str(e)
        print(f"Error processing document {document_id}: {error_message}")
        print(traceback.format_exc())
        
        # Update status to failed with error message (with org_id if available)
        try:
            # Try to get org_id from document for status update
            document = common.find_one(
                table='kb_docs',
                filters={'id': document_id}
            )
            if document and document.get('org_id'):
                update_document_status(document_id, document['org_id'], 'failed', error_message)
            else:
                # Fallback without org_id filter (shouldn't happen but handle gracefully)
                common.update_one(
                    table='kb_docs',
                    filters={'id': document_id},
                    data={'status': 'failed', 'error_message': error_message}
                )
        except Exception as update_error:
            print(f"Error updating failed status: {str(update_error)}")


# ============================================================================
# Document Parsing
# ============================================================================

def download_from_s3(bucket: str, key: str) -> Tuple[bytes, str]:
    """Download document from S3."""
    try:
        response = s3_client.get_object(Bucket=bucket, Key=key)
        document_bytes = response['Body'].read()
        mime_type = response.get('ContentType', 'application/octet-stream')
        
        return document_bytes, mime_type
    
    except Exception as e:
        raise ValueError(f"Failed to download from S3: {str(e)}")


def parse_document(document_bytes: bytes, mime_type: str, filename: str) -> Tuple[str, Dict]:
    """Parse document based on MIME type."""
    try:
        if mime_type == 'application/pdf' or filename.lower().endswith('.pdf'):
            return parse_pdf(document_bytes)
        
        elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                          'application/msword'] or filename.lower().endswith(('.docx', '.doc')):
            return parse_docx(document_bytes)
        
        elif mime_type in ['text/plain', 'text/markdown'] or filename.lower().endswith(('.txt', '.md')):
            return parse_text(document_bytes)
        
        else:
            raise ValueError(f"Unsupported document type: {mime_type}")
    
    except Exception as e:
        raise ValueError(f"Failed to parse document: {str(e)}")


def parse_pdf(document_bytes: bytes) -> Tuple[str, Dict]:
    """Parse PDF document using PyPDF2."""
    try:
        import PyPDF2
        
        pdf_file = BytesIO(document_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        page_count = len(pdf_reader.pages)
        
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        text = "\n\n".join(text_parts)
        
        # Extract metadata
        metadata = {
            'pageCount': page_count,
            'wordCount': len(text.split()),
            'parser': 'PyPDF2'
        }
        
        # Try to extract PDF metadata
        if pdf_reader.metadata:
            if pdf_reader.metadata.get('/Title'):
                metadata['title'] = str(pdf_reader.metadata.get('/Title'))
            if pdf_reader.metadata.get('/Author'):
                metadata['author'] = str(pdf_reader.metadata.get('/Author'))
            if pdf_reader.metadata.get('/CreationDate'):
                metadata['createdDate'] = str(pdf_reader.metadata.get('/CreationDate'))
        
        return text, metadata
    
    except ImportError:
        raise ValueError("PyPDF2 not available - install with: pip install PyPDF2")
    except Exception as e:
        raise ValueError(f"PDF parsing error: {str(e)}")


def parse_docx(document_bytes: bytes) -> Tuple[str, Dict]:
    """Parse DOCX document using python-docx."""
    try:
        import docx
        
        docx_file = BytesIO(document_bytes)
        doc = docx.Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        text = "\n\n".join(text_parts)
        
        # Extract metadata
        metadata = {
            'paragraphCount': len(doc.paragraphs),
            'wordCount': len(text.split()),
            'parser': 'python-docx'
        }
        
        # Try to extract DOCX core properties
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.created:
                metadata['createdDate'] = core_props.created.isoformat()
        except:
            pass
        
        return text, metadata
    
    except ImportError:
        raise ValueError("python-docx not available - install with: pip install python-docx")
    except Exception as e:
        raise ValueError(f"DOCX parsing error: {str(e)}")


def parse_text(document_bytes: bytes) -> Tuple[str, Dict]:
    """Parse plain text or markdown document."""
    try:
        # Try UTF-8 first, fallback to latin-1
        try:
            text = document_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text = document_bytes.decode('latin-1')
        
        metadata = {
            'wordCount': len(text.split()),
            'lineCount': len(text.splitlines()),
            'parser': 'text'
        }
        
        return text, metadata
    
    except Exception as e:
        raise ValueError(f"Text parsing error: {str(e)}")


# ============================================================================
# Text Chunking
# ============================================================================

def chunk_text(text: str, chunk_size: int, overlap: int) -> List[Dict[str, Any]]:
    """
    Split text into overlapping chunks, preserving sentence boundaries.
    
    Returns list of dicts with:
    - content: chunk text
    - chunk_index: position in document
    - metadata: {startChar, endChar}
    """
    chunks = []
    
    # Split into sentences (simple approach)
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    current_chunk = []
    current_length = 0
    start_char = 0
    chunk_index = 0
    
    for sentence in sentences:
        sentence_length = len(sentence)
        
        # If adding this sentence exceeds chunk_size and we have content
        if current_length + sentence_length > chunk_size and current_chunk:
            # Create chunk
            chunk_text = ' '.join(current_chunk)
            end_char = start_char + len(chunk_text)
            
            chunks.append({
                'content': chunk_text,
                'chunk_index': chunk_index,
                'metadata': {
                    'startChar': start_char,
                    'endChar': end_char,
                    'sentenceCount': len(current_chunk)
                }
            })
            
            chunk_index += 1
            
            # Calculate overlap: keep last N characters worth of sentences
            overlap_length = 0
            overlap_sentences = []
            
            for sent in reversed(current_chunk):
                if overlap_length + len(sent) <= overlap:
                    overlap_sentences.insert(0, sent)
                    overlap_length += len(sent)
                else:
                    break
            
            # Start new chunk with overlap
            current_chunk = overlap_sentences + [sentence]
            current_length = sum(len(s) for s in current_chunk)
            start_char = end_char - overlap_length
        
        else:
            current_chunk.append(sentence)
            current_length += sentence_length
    
    # Add final chunk if any content remains
    if current_chunk:
        chunk_text = ' '.join(current_chunk)
        chunks.append({
            'content': chunk_text,
            'chunk_index': chunk_index,
            'metadata': {
                'startChar': start_char,
                'endChar': start_char + len(chunk_text),
                'sentenceCount': len(current_chunk)
            }
        })
    
    return chunks


# ============================================================================
# Embedding Generation
# ============================================================================

def get_embedding_config() -> Dict[str, Any]:
    """
    Get embedding configuration from ai_cfg_sys_rag + ai_cfg_models tables.
    
    Uses foreign key relation to look up model details.
    Falls back to defaults if config not found or on error.
    """
    try:
        # Step 1: Get system-level RAG configuration
        sys_config = common.find_one(
            table='ai_cfg_sys_rag',
            filters={}  # Only one row exists (sys-level config)
        )
        
        if not sys_config:
            print("WARNING: No ai_cfg_sys_rag config found, using hardcoded defaults")
            return _default_embedding_config()
        
        # Step 2: Get embedding model ID
        embedding_model_id = sys_config.get('default_embedding_model_id')
        
        if not embedding_model_id:
            print("WARNING: No default_embedding_model_id in ai_cfg_sys_rag, using defaults")
            return _default_embedding_config()
        
        # Step 3: Look up model details in ai_models
        model_config = common.find_one(
            table='ai_models',
            filters={'id': embedding_model_id}
        )
        
        if not model_config:
            print(f"ERROR: Model {embedding_model_id} not found in ai_models, using defaults")
            return _default_embedding_config()
        
        # Step 4: Extract model details
        model_id = model_config.get('model_id')  # e.g., "amazon.titan-embed-text-v2:0"
        
        # Parse capabilities JSON to get embedding dimensions
        capabilities = model_config.get('capabilities', {})
        if isinstance(capabilities, str):
            import json
            capabilities = json.loads(capabilities)
        
        embedding_dimension = capabilities.get('embeddingDimensions', DEFAULT_EMBEDDING_DIMENSION)
        
        # Step 5: Get chunking parameters
        max_chunk_tokens = sys_config.get('max_chunk_size_tokens', 500)
        min_chunk_tokens = sys_config.get('min_chunk_size_tokens', 100)
        
        # Convert tokens to characters (rough estimate: 1 token ≈ 4 chars)
        chunk_size = max_chunk_tokens * 4
        chunk_overlap = min_chunk_tokens * 4  # Use min as overlap
        
        print(f"Loaded config from ai_cfg_sys_rag + ai_models:")
        print(f"  - Model ID: {embedding_model_id}")
        print(f"  - Model: {model_id}")
        print(f"  - Embedding dimensions: {embedding_dimension}")
        print(f"  - Chunk size: {chunk_size} chars ({max_chunk_tokens} tokens)")
        print(f"  - Chunk overlap: {chunk_overlap} chars ({min_chunk_tokens} tokens)")
        
        return {
            'provider': 'bedrock',  # Inferred from model_id prefix
            'model': model_id,
            'dimension': embedding_dimension,
            'chunkSize': chunk_size,
            'chunkOverlap': chunk_overlap
        }
    
    except Exception as e:
        print(f"ERROR loading embedding config: {str(e)}")
        print(traceback.format_exc())
        print("Falling back to hardcoded defaults")
        return _default_embedding_config()


def _default_embedding_config() -> Dict[str, Any]:
    """Return default embedding configuration."""
    return {
        'provider': 'bedrock',
        'model': DEFAULT_EMBEDDING_MODEL,
        'dimension': DEFAULT_EMBEDDING_DIMENSION,
        'chunkSize': DEFAULT_CHUNK_SIZE,
        'chunkOverlap': DEFAULT_CHUNK_OVERLAP
    }


def generate_embeddings(chunks: List[Dict], model: str) -> List[List[float]]:
    """
    Generate embeddings for all chunks using configured model.
    
    Supports batching for efficiency.
    """
    try:
        all_embeddings = []
        
        # Process in batches
        for i in range(0, len(chunks), BATCH_SIZE):
            batch_chunks = chunks[i:i + BATCH_SIZE]
            batch_embeddings = generate_embeddings_batch(batch_chunks, model)
            all_embeddings.extend(batch_embeddings)
        
        return all_embeddings
    
    except Exception as e:
        raise ValueError(f"Embedding generation failed: {str(e)}")


def generate_embeddings_batch(chunks: List[Dict], model: str) -> List[List[float]]:
    """Generate embeddings for a batch of chunks."""
    try:
        # AWS Bedrock Titan Text Embeddings V2
        if 'titan-embed' in model:
            return generate_bedrock_embeddings(chunks, model)
        
        # OpenAI (future support)
        elif 'text-embedding' in model:
            raise NotImplementedError("OpenAI embeddings not yet implemented")
        
        # Unknown provider
        else:
            raise ValueError(f"Unsupported embedding model: {model}")
    
    except Exception as e:
        raise ValueError(f"Batch embedding error: {str(e)}")


def generate_bedrock_embeddings(chunks: List[Dict], model: str) -> List[List[float]]:
    """Generate embeddings using AWS Bedrock Titan."""
    try:
        embeddings = []
        
        for chunk in chunks:
            # Bedrock API call
            request_body = {
                'inputText': chunk['content']
            }
            
            response = bedrock_runtime.invoke_model(
                modelId=model,
                contentType='application/json',
                accept='application/json',
                body=json.dumps(request_body)
            )
            
            response_body = json.loads(response['body'].read())
            embedding = response_body.get('embedding')
            
            if not embedding:
                raise ValueError("No embedding in Bedrock response")
            
            embeddings.append(embedding)
        
        return embeddings
    
    except Exception as e:
        raise ValueError(f"Bedrock embedding error: {str(e)}")


# ============================================================================
# Database Operations
# ============================================================================

def get_kb_config(kb_id: str, org_id: str) -> Dict[str, Any]:
    """Get KB configuration from database with org_id filter (CORA Compliance)."""
    try:
        kb = common.find_one(
            table='kb_bases',
            filters={'id': kb_id, 'org_id': org_id, 'is_deleted': False}
        )
        
        if not kb:
            raise ValueError(f"KB {kb_id} not found or access denied")
        
        return kb.get('config') or {}
    
    except Exception as e:
        print(f"Error getting KB config: {str(e)}")
        raise


def store_chunks(document_id: str, kb_id: str, org_id: str, chunks: List[Dict], 
                embeddings: List[List[float]], embedding_model: str):
    """Store chunks with embeddings in database (with org_id for multi-tenancy)."""
    try:
        for chunk, embedding in zip(chunks, embeddings):
            # Convert embedding list to pgvector format string
            # pgvector accepts '[x,y,z]' format for vector columns
            embedding_str = '[' + ','.join(map(str, embedding)) + ']'
            
            # Use common.insert_one for inserts
            # The database driver handles the vector type conversion
            common.insert_one(
                table='kb_chunks',
                data={
                    'kb_id': kb_id,
                    'document_id': document_id,
                    'content': chunk['content'],
                    'embedding': embedding_str,
                    'chunk_index': chunk['chunk_index'],
                    'token_count': estimate_token_count(chunk['content']),
                    'metadata': chunk['metadata'],
                    'embedding_model': embedding_model,
                    'org_id': org_id
                }
            )
        
        print(f"Stored {len(chunks)} chunks for document {document_id}")
    
    except Exception as e:
        raise ValueError(f"Failed to store chunks: {str(e)}")


def update_document_status(document_id: str, org_id: str, status: str, error_message: Optional[str] = None):
    """Update document processing status with org_id filter (CORA Compliance)."""
    try:
        update_data = {'status': status}
        
        if error_message:
            update_data['error_message'] = error_message
        else:
            update_data['error_message'] = None
        
        common.update_one(
            table='kb_docs',
            filters={'id': document_id, 'org_id': org_id},
            data=update_data
        )
        
        print(f"Updated document {document_id} status to: {status}")
    
    except Exception as e:
        print(f"Error updating document status: {str(e)}")
        # Don't raise - status update failure shouldn't fail processing


def update_document_metadata(document_id: str, org_id: str, metadata: Dict, chunk_count: int):
    """Update document metadata after successful processing with org_id filter (CORA Compliance)."""
    try:
        common.update_one(
            table='kb_docs',
            filters={'id': document_id, 'org_id': org_id},
            data={
                'metadata': metadata,
                'chunk_count': chunk_count
            }
        )
    
    except Exception as e:
        print(f"Error updating document metadata: {str(e)}")
        # Don't raise - metadata update failure shouldn't fail processing


# ============================================================================
# Helper Functions
# ============================================================================

def estimate_token_count(text: str) -> int:
    """
    Estimate token count for billing tracking.
    
    Rough estimate: 1 token ≈ 4 characters for English text.
    """
    return len(text) // 4


def clean_text(text: str) -> str:
    """Clean extracted text."""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters that cause issues
    text = text.replace('\x00', '')
    
    return text.strip()
